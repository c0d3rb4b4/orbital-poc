"""Build the Google Docs-compatible case-study DOCX from its Markdown source.

Run with a Python environment containing python-docx:
    python docs/build_case_study_docx.py
"""

from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Orbital-Customer-Account-POC-Case-Study.md"
OUTPUT = ROOT / "Orbital-Customer-Account-POC-Case-Study.docx"

NAVY = "17365D"
BLUE = "2F5597"
PALE_BLUE = "EAF2F8"
PALE_GREY = "F3F5F7"
MID_GREY = "667085"
WHITE = "FFFFFF"
AMBER = "E2A93B"
GREEN = "3B7A57"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, color="D8DEE8", size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_run_font(run, name="Arial", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_run_font(run, size=8, color=MID_GREY)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.15)
    section.right_margin = Cm(2.15)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor.from_string("263238")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Heading 1", 18, NAVY, 15, 7),
        ("Heading 2", 13.5, BLUE, 12, 5),
        ("Heading 3", 11.2, NAVY, 9, 3),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.page_break_before = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(9.4)
        style.paragraph_format.space_after = Pt(3)

    if "Case Caption" not in styles:
        caption = styles.add_style("Case Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Case Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    caption.font.size = Pt(8.4)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MID_GREY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    if "Code Inline" not in styles:
        code_style = styles.add_style("Code Inline", WD_STYLE_TYPE.CHARACTER)
    else:
        code_style = styles["Code Inline"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string("7A2E0E")

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    props = doc.core_properties
    props.title = "Orbital Customer Account Integration POC — Case Study"
    props.subject = "Learnings and comparison with the current Azure BIP implementation"
    props.author = "Customer Account Integration POC Team"
    props.keywords = "Orbital, Taxi, RabbitMQ, Azure, BIP, Adobe, SAP, FWT, case study"
    props.comments = "Generated from the checked-in Markdown source."
    props.created = datetime(2026, 7, 29)
    props.modified = datetime(2026, 7, 30)


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("ORBITAL CUSTOMER ACCOUNT POC  |  CASE STUDY")
    set_run_font(r, size=7.5, color=MID_GREY, bold=True)
    p.paragraph_format.space_after = Pt(0)
    p_bdr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), AMBER)
    borders.append(bottom)
    p_bdr.append(borders)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = fp.add_run("29 July 2026   •   Page ")
    set_run_font(rr, size=8, color=MID_GREY)
    add_field(fp, "PAGE")
    rr = fp.add_run(" of ")
    set_run_font(rr, size=8, color=MID_GREY)
    add_field(fp, "NUMPAGES")


def add_accent_rule(doc: Document, color=AMBER, width="24") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), width)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


INLINE_PATTERN = re.compile(r"(<br\s*/?>|`[^`]+`|\*\*[^*]+\*\*)", re.IGNORECASE)


def add_inline(paragraph, text: str, *, default_bold=False, default_italic=False) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, bold=default_bold, italic=default_italic)
        token = match.group(0)
        if token.lower().startswith("<br"):
            paragraph.add_run().add_break()
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.style = "Code Inline"
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True, italic=default_italic)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, bold=default_bold, italic=default_italic)


def add_callout(doc: Document, title: str, body: str, fill=PALE_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    content = table.cell(0, 0)
    set_cell_shading(content, fill)
    set_cell_margins(content, 160, 260, 150, 180)
    tc_pr = content._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "32")
    left.set(qn("w:color"), accent)
    borders.append(left)
    tc_pr.append(borders)
    p = content.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    title_run = p.add_run(title.upper())
    set_run_font(title_run, size=8.5, color=accent, bold=True)
    p = content.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(22)
    badge = doc.add_table(rows=1, cols=1)
    badge.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = badge.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, 95, 170, 95, 170)
    run = cell.paragraphs[0].add_run("CASE STUDY")
    set_run_font(run, size=9, color=WHITE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Orbital Customer Account\nIntegration POC")
    set_run_font(r, size=29, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("A case study in semantic integration with Taxi")
    set_run_font(r, size=16, color=BLUE, bold=True)

    add_accent_rule(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run(
        "Adobe, SAP and FWT customer-account updates using Orbital, Taxi, "
        "RabbitMQ and a narrow Python transport bridge"
    )
    set_run_font(r, size=11, color="344054")

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    for row, (label, value) in zip(
        meta.rows,
        (
            ("PREPARED", "29 July 2026"),
            ("DOCUMENT STATUS", "Implementation case study"),
            ("SCOPE", "Bounded proof of concept; not a production-readiness approval"),
        ),
    ):
        row.cells[0].width = Cm(3.3)
        row.cells[1].width = Cm(12.5)
        for c in row.cells:
            set_cell_margins(c, 75, 80, 75, 80)
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        set_run_font(lr, size=7.5, color=MID_GREY, bold=True)
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(value)
        set_run_font(vr, size=9.2, color=NAVY, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_callout(
        doc,
        "Decision position",
        "The POC proves that reusable Taxi semantics can drive a selected customer-account "
        "update slice across three independent wire contracts. It does not yet prove that "
        "Orbital and RabbitMQ should replace the current Azure integration platform. The "
        "recommended next step is a controlled production-readiness pilot, with the Taxi "
        "semantic layer retained as the main demonstrated value and broker replacement "
        "assessed separately.",
    )


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        for col_index in range(columns):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, PALE_GREY)
            text = values[col_index] if col_index < len(values) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if len(rows) <= 11 and row_index < len(rows) - 1:
                p.paragraph_format.keep_with_next = True
            add_inline(p, text, default_bold=row_index == 0)
            for run in p.runs:
                if row_index == 0:
                    set_run_font(run, size=8.2, color=WHITE, bold=True)
                else:
                    if run.style is None or run.style.name != "Code Inline":
                        set_run_font(run, size=8.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table_line(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_line(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def set_picture_alt_text(paragraph, description: str) -> None:
    for drawing in paragraph._p.xpath(".//wp:docPr"):
        drawing.set("descr", description)


def add_image(doc: Document, alt: str, path_text: str) -> None:
    path = (ROOT / path_text).resolve()
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.55))
    set_picture_alt_text(p, alt)


def add_body_paragraph(doc: Document, text: str) -> None:
    is_figure_caption = text.startswith("**Figure ")
    is_table_caption = text.startswith("**Table ")
    style = "Case Caption" if is_figure_caption else None
    p = doc.add_paragraph(style=style)
    if is_figure_caption:
        stripped = text.replace("**", "")
        add_inline(p, stripped, default_italic=True)
    elif is_table_caption:
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        add_inline(p, text)
    else:
        add_inline(p, text)


def parse_markdown(doc: Document, lines: list[str]) -> None:
    index = 0
    paragraph_buffer: list[str] = []
    suppress_next_heading_page_break = False

    def flush_paragraph() -> None:
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("<!--"):
            flush_paragraph()
            if stripped == "<!-- PAGE BREAK -->":
                doc.add_page_break()
            elif stripped == "<!-- CONTINUE PAGE -->":
                suppress_next_heading_page_break = True
            index += 1
            continue

        image_match = re.fullmatch(r"!\[(.+)]\((.+)\)", stripped)
        if image_match:
            flush_paragraph()
            add_image(doc, image_match.group(1), image_match.group(2))
            index += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            p = doc.add_heading(heading.group(2), level=level)
            if level == 1:
                p.paragraph_format.page_break_before = not suppress_next_heading_page_break
                p_pr = p._p.get_or_add_pPr()
                borders = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "14")
                bottom.set(qn("w:space"), "5")
                bottom.set(qn("w:color"), AMBER)
                borders.append(bottom)
                p_pr.append(borders)
            suppress_next_heading_page_break = False
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1].strip()):
            flush_paragraph()
            table_rows = [parse_table_line(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(parse_table_line(lines[index].strip()))
                index += 1
            add_table(doc, table_rows)
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip().lstrip(">").strip())
                index += 1
            text = " ".join(quote_lines)
            title = "Key point"
            if "  " in text:
                maybe_title, body = text.split("  ", 1)
                title = maybe_title.replace("**", "")
                text = body
            add_callout(doc, title, text.replace("**", ""))
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.55)
            p.paragraph_format.space_after = Pt(3)
            marker = p.add_run(f"{numbered.group(1)}.  ")
            set_run_font(marker, color=BLUE, bold=True)
            add_inline(p, numbered.group(2))
            index += 1
            continue

        if stripped == "---":
            flush_paragraph()
            add_accent_rule(doc, color="D8DEE8", width="8")
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def remove_consecutive_page_breaks(doc: Document) -> None:
    """Avoid empty pages when a manual break precedes a page-break-before heading."""
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs[:-1]):
        if paragraph._p.xpath(".//w:br[@w:type='page']"):
            next_p = paragraphs[idx + 1]
            if next_p.style and next_p.style.name == "Heading 1":
                next_p.paragraph_format.page_break_before = False


def build(output: Path = OUTPUT) -> Path:
    source_text = SOURCE.read_text(encoding="utf-8")
    lines = source_text.splitlines()
    try:
        body_start = lines.index("<!-- PAGE BREAK -->") + 1
    except ValueError as exc:
        raise ValueError("The case-study source is missing its cover page break marker") from exc

    doc = Document()
    configure_document(doc)
    configure_header_footer(doc)
    add_cover(doc)
    doc.add_page_break()
    parse_markdown(doc, lines[body_start:])
    remove_consecutive_page_breaks(doc)
    output = output.resolve()
    doc.save(output)
    return output


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--output",
        type=Path,
        default=OUTPUT,
        help=f"DOCX output path (default: {OUTPUT})",
    )
    args = parser.parse_args()
    result = build(args.output)
    print(result)
